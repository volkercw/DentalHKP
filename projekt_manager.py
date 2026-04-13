"""
DentalHKP – Projekt-Management
- Speichern der Analyse als JSON (D:\\DentalProjekte\\HKP\\...)
- Word-Dokument Generierung (mit Briefvorlage Dr. Jung)
- Projekt-Liste
"""
import json
import re
import datetime
from pathlib import Path
from decimal import Decimal
from config import PROJEKTE_PFAD


# ─────────────────────────────────────────────────────────────────────────────
# Hilfsfunktionen
# ─────────────────────────────────────────────────────────────────────────────

def _goz_display(goz_nr: str) -> str:
    """Entfernt Charly-interne Material-Suffixe für offizielle Ausgabe.

    Charly hängt Material-Kürzel an GOZ-Nummern:
      k = Keramik  →  2170k → 2170
      v = Verblendet → 2210v → 2210
      z = Zirkon    → 2120z → 2120
      a = Akryl/Analog → 5190a → 5190

    Behalten wird 'i' (§6-Analog Implantat: 2200i bleibt 2200i,
    da es intern als eigener Preisschlüssel geführt wird).
    Sonderzeichen wie 'Ä1' (Ärztliche Beratung) bleiben unverändert.
    """
    if not goz_nr:
        return goz_nr
    return re.sub(r'^(\d+)[kvza]$', r'\1', goz_nr)


def _safe(obj):
    """Decimal/None-safe Konvertierung für JSON."""
    if isinstance(obj, Decimal):
        return float(obj)
    if isinstance(obj, dict):
        return {k: _safe(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_safe(v) for v in obj]
    return obj


# ─────────────────────────────────────────────────────────────────────────────
# Projekt-Daten aufbauen
# ─────────────────────────────────────────────────────────────────────────────

def build_hkp_projekt(
    kv_details: dict,
    patient_info: dict,
    gap_teeth: list[dict],
    agent_result: dict,
    manual_goz: list[dict] = None,
    kv_material: list[dict] = None,
    kv_labor: dict = None,
    approved_by: str = None,
    selected_positions: list[dict] = None,
    goz_prices: dict = None,
) -> dict:
    """
    Baut das vollständige HKP-Projekt-JSON auf.
    Enthält alle Daten für spätere Weiterarbeit und Word-Generierung.
    """
    now = datetime.datetime.now()

    positionen = kv_details.get("positionen", []) if kv_details else []
    goz_entries = []
    for p in positionen:
        if p.get("goz_nr"):
            goz_entries.append({
                "goz_nr": p.get("goz_nr"),
                "goz_text": p.get("goz_text", ""),
                "faktor": float(p.get("faktor") or 0),
                "anzahl": float(p.get("anzahl") or 1),
                "betrag": float(p.get("betrag") or 0),
                "zahn_bitmask": p.get("zahn_bitmask"),
                "phase": p.get("phase_bezeichnung", ""),
            })

    return {
        "schema_version": "2.0",
        "erstellt_am": now.isoformat(),
        "genehmigt_von": approved_by,
        "charly_eingespielt": False,

        "patient": {
            "id": patient_info.get("solid"),
            "name": patient_info.get("name", ""),
            "vorname": patient_info.get("vorname", ""),
            "geburtsdatum": patient_info.get("geburtsdatum", ""),
        },
        "kv": {
            "id": kv_details.get("solid") if kv_details else None,
            "kurztext": kv_details.get("kurztext", "") if kv_details else "",
            "datum": kv_details.get("datum", "") if kv_details else "",
            "status": kv_details.get("kvstatus") if kv_details else None,
            "honorar": float(kv_details.get("honorar") or 0) if kv_details else 0,
            "material": float(kv_details.get("material") or 0) if kv_details else 0,
            "labor": float(kv_details.get("labor") or 0) if kv_details else 0,
        },

        "gap_teeth": _safe(gap_teeth),
        "bestehende_goz": _safe(goz_entries),

        "agent_ergebnisse": {
            "archiv_analyse": agent_result.get("archiv", ""),
            "goz_vorschlaege": _safe(agent_result.get("goz_structured", {})),
            "qualitaetspruefung": agent_result.get("qualitaet", ""),
        },

        "manual_goz": _safe(manual_goz or []),
        "material": _safe(kv_material or []),
        "labor": _safe(kv_labor or {}),

        # Eigenständiger Kostenvoranschlag aus KI-Analyse
        "kostenvoranschlag": _build_kva(selected_positions or [], goz_prices or {}),
    }


def _group_kva_positions(positionen: list[dict]) -> list[dict]:
    """
    Fasst identische GOZ-Positionen zusammen (gleiche goz_nr + faktor).
    Zähne werden zu einer kommagetrennten Liste, Anzahl/Beträge summiert.
    Reihenfolge bleibt erhalten (erste Erwähnung bestimmt Position).
    """
    from collections import OrderedDict
    groups: dict = OrderedDict()
    for pos in positionen:
        key = (pos["goz_nr"], float(pos.get("faktor", 2.3)))
        if key not in groups:
            groups[key] = {
                "zaehne":  [],
                "goz_nr":  pos["goz_nr"],
                "text":    pos.get("text", ""),
                "faktor":  float(pos.get("faktor", 2.3)),
                "anzahl":  0,
                "honorar": 0.0,
                "mat_est": 0.0,
                "lab_est": 0.0,
                "farbe":   pos.get("farbe", "none"),
            }
        groups[key]["zaehne"].append(str(pos["zahn"]))
        groups[key]["anzahl"]  += int(pos.get("anzahl", 1))
        groups[key]["honorar"] += float(pos.get("honorar", 0))
        groups[key]["mat_est"] += float(pos.get("mat_est", 0))
        groups[key]["lab_est"] += float(pos.get("lab_est", 0))

    from config import GOZ_SESSION_EINMALIG
    result = []
    for g in groups.values():
        g["zahn_str"] = ", ".join(g.pop("zaehne"))
        # Session-GOZ: Anzahl auf 1 begrenzen + Honorar anteilig kürzen
        if g["goz_nr"] in GOZ_SESSION_EINMALIG and g["anzahl"] > 1:
            factor = 1.0 / g["anzahl"]
            g["honorar"] = round(g["honorar"] * factor, 2)
            g["mat_est"] = round(g["mat_est"] * factor, 2)
            g["lab_est"] = round(g["lab_est"] * factor, 2)
            g["anzahl"] = 1
            g["zahn_str"] = g["zahn_str"].split(",")[0].strip() + " (Sitzung)"
        result.append(g)
    return result


def _build_kva(selected_positions: list[dict], goz_prices: dict) -> dict:
    """
    Berechnet Kostenvoranschlag-Totals aus gewählten GOZ-Positionen.
    Session-GOZ (MKO-Paket etc.) werden automatisch dedupliziert –
    nur die erste Nennung zählt, egal wie oft der Arzt sie für mehrere
    Zähne ausgewählt hat.
    """
    from config import GOZ_SESSION_EINMALIG
    positionen_out = []
    total_h = total_m = total_l = 0.0
    seen_session_goz: set[str] = set()   # Sicherheitsnetz gegen Doppel-MKO

    for r in selected_positions:
        goz_nr = r.get("goz_nr", "")

        # Session-GOZ: nur einmal aufnehmen (safety net falls Agent dennoch doppelt)
        if goz_nr in GOZ_SESSION_EINMALIG:
            if goz_nr in seen_session_goz:
                continue          # überspringe Duplikat
            seen_session_goz.add(goz_nr)

        p = goz_prices.get(goz_nr, {})
        base = float(p.get("base_fee", 0))
        hon  = base * float(r.get("faktor", 2.3)) * int(r.get("anzahl", 1))
        mat  = float(p.get("mat_est", 0)) * int(r.get("anzahl", 1))
        lab  = float(p.get("lab_est", 0)) * int(r.get("anzahl", 1))
        total_h += hon; total_m += mat; total_l += lab
        positionen_out.append({
            "zahn":    r.get("zahn"),
            "goz_nr":  goz_nr,
            "text":    r.get("text", ""),
            "faktor":  float(r.get("faktor", 2.3)),
            "anzahl":  int(r.get("anzahl", 1)),
            "farbe":   r.get("farbe", "none"),
            "honorar": round(hon, 2),
            "mat_est": round(mat, 2),
            "lab_est": round(lab, 2),
        })
    labor_brutto = total_l * 1.07  # 7% MwSt Fremdlabor
    return {
        "positionen": positionen_out,
        "honorar":      round(total_h, 2),
        "material_est": round(total_m, 2),
        "labor_est":    round(total_l, 2),
        "labor_brutto": round(labor_brutto, 2),
        "gesamt":       round(total_h + total_m + labor_brutto, 2),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Speichern
# ─────────────────────────────────────────────────────────────────────────────

def save_projekt(projekt: dict) -> Path:
    """Speichert Projekt in D:\\DentalProjekte\\HKP\\{datum}_{name}_KV{id}\\"""
    now = datetime.datetime.now()
    name = f"{projekt['patient']['name']}_{projekt['patient']['vorname']}"
    name = name.replace(" ", "_").replace(",", "")
    kv_id = projekt["kv"]["id"] or "?"

    ordner_name = f"{now.strftime('%Y-%m-%d_%H%M')}_HKP_{name}_KV{kv_id}"
    ordner = Path(PROJEKTE_PFAD) / "HKP" / ordner_name
    ordner.mkdir(parents=True, exist_ok=True)

    with open(ordner / "projekt_daten.json", "w", encoding="utf-8") as f:
        json.dump(projekt, f, ensure_ascii=False, indent=2)

    meta = {
        "erstellt_am": projekt["erstellt_am"],
        "patient": f"{projekt['patient']['name']}, {projekt['patient']['vorname']}",
        "kv_id": kv_id,
        "datum": projekt["kv"]["datum"],
        "honorar": projekt["kv"]["honorar"],
        "labor": projekt["kv"]["labor"],
        "gap_anzahl": len(projekt["gap_teeth"]),
        "charly_eingespielt": projekt["charly_eingespielt"],
    }
    with open(ordner / "projekt_meta.json", "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    return ordner


# ─────────────────────────────────────────────────────────────────────────────
# Word-Dokument generieren
# ─────────────────────────────────────────────────────────────────────────────

def generate_word(projekt: dict, output_path: Path):
    """Generiert Word-Dokument mit Briefvorlage Dr. Jung."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from collections import defaultdict
    import re

    template_path = Path(__file__).parent / "Dr. Jung Zahnklinik Briefvorlage.dotx"
    if template_path.exists():
        # .dotx → .docx: ContentType-Patch (wie in DentalAI) damit Header/Logo erhalten bleibt
        import zipfile, shutil, tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        shutil.copy2(str(template_path), tmp_path)
        # ContentType patchen: alle Entries neu schreiben (kein Duplicate)
        with zipfile.ZipFile(tmp_path, 'r') as zin:
            names = zin.namelist()
            entries = {n: zin.read(n) for n in names}
        ct = entries['[Content_Types].xml'].decode('utf-8')
        entries['[Content_Types].xml'] = ct.replace(
            'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
        ).encode('utf-8')
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                zout.writestr(name, entries[name])
        doc = Document(tmp_path)
        Path(tmp_path).unlink(missing_ok=True)
        # Body leeren – sectPr (Seitenlayout) ZUERST sichern
        from docx.oxml.ns import qn as _qn
        body = doc.element.body
        sect_pr = body.find(_qn('w:sectPr'))
        for elem in list(body):
            body.remove(elem)
        # sectPr wieder einfügen (sonst fehlt das Seitenlayout)
        if sect_pr is not None:
            body.append(sect_pr)
        # Seitenränder wie DentalAI (top:4cm, bottom:2cm, left:2.5cm, right:4cm)
        from docx.shared import Cm
        for section in doc.sections:
            section.top_margin    = Cm(4)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(4)
    else:
        doc = Document()

    # ── Hilfsfunktionen ──────────────────────────────────────────────────────

    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def add_heading(text: str, level: int = 1):
        """Template-kompatible Überschrift via Normal-Style + Bold + Größe."""
        sizes = {1: 18, 2: 14, 3: 12, 4: 11}
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 11))
        if level <= 2:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
        return p

    def strip_md(text: str) -> str:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        return re.sub(r'^#{1,4}\s+', '', text, flags=re.MULTILINE).strip()

    def add_markdown_para(text: str):
        for line in text.split('\n'):
            if not line.strip():
                doc.add_paragraph("", style='Normal')
                continue
            if re.match(r'^#{1,4}\s', line):
                lvl = min(len(re.match(r'^(#+)', line).group(1)) + 1, 4)
                add_heading(line.lstrip('#').strip(), level=lvl)
            elif line.startswith(('- ', '* ')):
                doc.add_paragraph(line[2:], style='List Paragraph')
            else:
                p = doc.add_paragraph(style='Normal')
                # Inline-Bold (**text**)
                parts = re.split(r'\*\*(.+?)\*\*', strip_md(line))
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.bold = (i % 2 == 1)

    farbe_label = {"gruen": "PFLICHT", "gelb": "EMPFOHLEN"}
    farbe_hex   = {"gruen": "D4EDDA", "gelb": "FFF3CD"}

    # ── Titel ────────────────────────────────────────────────────────────────
    pat = projekt["patient"]
    kv  = projekt["kv"]

    h = add_heading("Heil- und Kostenplan – GOZ-Analyse", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Patient: {pat['name']}, {pat['vorname']}  |  KV #{kv['id']} vom {kv['datum']}  |  "
        f"Erstellt: {projekt['erstellt_am'][:10]}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if projekt.get("genehmigt_von"):
        p = doc.add_paragraph(f"Freigabe: {projekt['genehmigt_von']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── KV-Übersicht ─────────────────────────────────────────────────────────
    add_heading("KV-Übersicht", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(["Bezeichnung", "Honorar", "Material", "Labor"]):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    r = t.add_row().cells
    r[0].text = kv.get("kurztext", "–")
    r[1].text = f"{kv.get('honorar', 0):.2f} €"
    r[2].text = f"{kv.get('material', 0):.2f} €"
    r[3].text = f"{kv.get('labor', 0):.2f} €"
    doc.add_paragraph("")

    # ── GOZ-Lücken ───────────────────────────────────────────────────────────
    gap_teeth = projekt.get("gap_teeth", [])
    if gap_teeth:
        add_heading("GOZ-Lücken (grafisch geplant, ohne Abrechnung)", level=2)
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Table Grid'
        for i, h in enumerate(["Zahn (FDI)", "Behandlung", "GOZ-Basis"]):
            t2.rows[0].cells[i].text = h
            t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for tooth in gap_teeth:
            r = t2.add_row().cells
            r[0].text = str(tooth.get("zahn", "?"))
            r[1].text = tooth.get("treatment_name", "?")
            r[2].text = tooth.get("goz_basis", "?") or "–"
        doc.add_paragraph("")

    # ── GOZ-Vorschläge ───────────────────────────────────────────────────────
    goz_struct = projekt.get("agent_ergebnisse", {}).get("goz_vorschlaege", {})
    manual_goz = projekt.get("manual_goz", [])

    add_heading("KI-Vorschläge: GOZ-Positionen", level=2)
    doc.add_paragraph(
        "🟢 Pflicht (grün)  |  🟡 Empfohlen (gelb)  |  ⬜ Optional  |  🟣 Manuell hinzugefügt"
    )

    gesamtbeg = goz_struct.get("gesamtbegruendung", "")
    if gesamtbeg:
        doc.add_paragraph(gesamtbeg)

    for zahn_data in goz_struct.get("zaehne", []):
        zahn_nr    = zahn_data.get("zahn", "?")
        behandlung = zahn_data.get("behandlung", "")
        alle_pos   = list(zahn_data.get("positionen", []))
        for m in manual_goz:
            if m.get("zahn") == zahn_nr:
                alle_pos.append({**m, "_manuell": True})

        add_heading(f"Zahn {zahn_nr} – {behandlung}", level=3)

        if not alle_pos:
            doc.add_paragraph("Keine Positionen vorgeschlagen.")
            continue

        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        for i, h in enumerate(["GOZ-Nr.", "Leistung", "Faktor", "Anz.", "Status"]):
            c = tbl.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True

        for pos in alle_pos:
            farbe  = pos.get("farbe")
            is_man = pos.get("_manuell", False)
            status = "Manuell" if is_man else farbe_label.get(farbe, "Optional")
            bg     = "E8D5F5" if is_man else farbe_hex.get(farbe, "F8F9FA")

            row = tbl.add_row()
            vals = [
                _goz_display(pos.get("goz_nr", "")),
                (pos.get("text") or pos.get("bezeichnung") or "")[:60],
                str(pos.get("faktor", "")),
                str(pos.get("anzahl", "")),
                status,
            ]
            for i, (cell, val) in enumerate(zip(row.cells, vals)):
                cell.text = val
                set_cell_bg(cell, bg)
        doc.add_paragraph("")

    # ── Archiv-Analyse ───────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading("Archiv-Analyse (historische Muster)", level=2)
    archiv = projekt.get("agent_ergebnisse", {}).get("archiv_analyse", "")
    add_markdown_para(archiv) if archiv else doc.add_paragraph("Keine Daten.")

    # ── Qualitätsprüfung ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading("Qualitätsprüfung", level=2)
    qualitaet = projekt.get("agent_ergebnisse", {}).get("qualitaetspruefung", "")
    add_markdown_para(qualitaet) if qualitaet else doc.add_paragraph("Keine Daten.")

    # ── Material ─────────────────────────────────────────────────────────────
    mat_data = projekt.get("material", [])
    if mat_data:
        doc.add_page_break()
        add_heading("Material", level=2)
        by_phase = defaultdict(list)
        for m in mat_data:
            by_phase[m.get("phase", "–")].append(m)
        mat_total = sum(float(m.get("betrag") or 0) for m in mat_data)
        doc.add_paragraph(f"Gesamt: {mat_total:.2f} €")
        for phase, items in by_phase.items():
            add_heading(phase, level=3)
            tbl_m = doc.add_table(rows=1, cols=4)
            tbl_m.style = 'Table Grid'
            for i, h in enumerate(["Kürzel", "Bezeichnung", "Anzahl", "Betrag"]):
                tbl_m.rows[0].cells[i].text = h
                tbl_m.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for item in items:
                r = tbl_m.add_row().cells
                r[0].text = str(item.get("mat_kuerzel") or item.get("mat_nr") or "")
                r[1].text = str(item.get("mat_bez") or "")[:60]
                r[2].text = str(item.get("anzahl") or "")
                r[3].text = f"{float(item.get('betrag') or 0):.2f} €"
            doc.add_paragraph("")

    # ── Labor ────────────────────────────────────────────────────────────────
    labor_data   = projekt.get("labor", {})
    lab_lsts     = labor_data.get("leistungen", [])
    lab_summary  = labor_data.get("summary", {})
    lab_mat      = labor_data.get("materialien", [])

    if lab_lsts or lab_summary:
        doc.add_page_break()
        add_heading("Labor", level=2)
        if lab_summary:
            fremd     = float(lab_summary.get("fremd_labor") or 0)
            fremd_mat = float(lab_summary.get("fremd_material") or 0)
            doc.add_paragraph(
                f"Fremdlabor: {fremd:.2f} €  |  Labor-Material: {fremd_mat:.2f} €  |  "
                f"Gesamt: {fremd + fremd_mat:.2f} €"
            )
        if lab_lsts:
            add_heading("Labor-Leistungen", level=3)
            tbl_l = doc.add_table(rows=1, cols=4)
            tbl_l.style = 'Table Grid'
            for i, h in enumerate(["Nr.", "Bezeichnung", "Anzahl", "Betrag"]):
                tbl_l.rows[0].cells[i].text = h
                tbl_l.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for lp in lab_lsts:
                r = tbl_l.add_row().cells
                r[0].text = str(lp.get("nummer") or "")
                r[1].text = str(lp.get("bezeichnung") or "")[:60]
                r[2].text = str(lp.get("anzahl") or "")
                r[3].text = f"{float(lp.get('betrag') or 0):.2f} €"
        if lab_mat:
            add_heading("Labor-Material", level=3)
            tbl_lm = doc.add_table(rows=1, cols=4)
            tbl_lm.style = 'Table Grid'
            for i, h in enumerate(["Kürzel", "Bezeichnung", "Anzahl", "Betrag"]):
                tbl_lm.rows[0].cells[i].text = h
                tbl_lm.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for lm in lab_mat:
                r = tbl_lm.add_row().cells
                r[0].text = str(lm.get("kuerzel") or lm.get("nummer") or "")
                r[1].text = str(lm.get("bezeichnung") or "")[:60]
                r[2].text = str(lm.get("anzahl") or "")
                r[3].text = f"{float(lm.get('betrag') or 0):.2f} €"

    doc.save(str(output_path))


# ─────────────────────────────────────────────────────────────────────────────
# Angebots-Dokument (nach PDF-Vorlage: Anschreiben + Anlage 1/2/3)
# ─────────────────────────────────────────────────────────────────────────────

def generate_angebot_word(
    kv_details: dict,
    patient_info: dict,
    kv_material: list[dict],
    kv_labor: dict,
    output_path: Path,
    selected_positions: list[dict] = None,
    goz_prices: dict = None,
    gruppiert: bool = True,
):
    """
    Erstellt ein Angebots-Word nach PDF-Vorlage:
      Seite 1  : Anschreiben + Kostenzusammenfassung
      Anlage 1 : Kostenvoranschlag Honorar (GOZ-Tabelle)
      Anlage 2 : Kostenvoranschlag Material
      Anlage 3 : Kostenvoranschlag Labor
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import zipfile, shutil, tempfile, datetime

    template_path = Path(__file__).parent / "Dr. Jung Zahnklinik Briefvorlage.dotx"
    if template_path.exists():
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        shutil.copy2(str(template_path), tmp_path)
        with zipfile.ZipFile(tmp_path, 'r') as zin:
            names = zin.namelist()
            entries = {n: zin.read(n) for n in names}
        entries['[Content_Types].xml'] = entries['[Content_Types].xml'].decode().replace(
            'template.main+xml', 'document.main+xml').encode()
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for n in names: zout.writestr(n, entries[n])
        doc = Document(tmp_path)
        Path(tmp_path).unlink(missing_ok=True)
        body = doc.element.body
        sect_pr = body.find(qn('w:sectPr'))
        for elem in list(body): body.remove(elem)
        if sect_pr is not None: body.append(sect_pr)
        for section in doc.sections:
            section.top_margin    = Cm(4)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(3.5)
    else:
        doc = Document()

    # ── Hilfsfunktionen ──────────────────────────────────────────────────────

    def set_cell_bg(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def para(text="", bold=False, size=10, align=None, style="Normal", space_before=0, space_after=4):
        p = doc.add_paragraph(style=style)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if align: p.alignment = align
        return p

    def section_header(patient_name: str, anlage: str, seite_text: str, titel: str):
        """Kopfzeile für Anlage-Seiten wie im PDF."""
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"Name: {patient_name}")
        r.font.size = Pt(9)
        # Tabellenzeile als Header
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        c = t.rows[0].cells
        # add_run() statt cell.text = ... um sicherzustellen dass immer ein Run existiert
        for cell, text in zip(c, [f"Name: {patient_name}", anlage or "", titel or ""]):
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(9)
            run.bold = True
        return t

    def fdi_from_pos(pos: dict) -> str:
        """FDI-Zahnbezeichnung aus fuellungszahn oder zahn_bitmask."""
        fz = pos.get("fuellungszahn")
        if fz and int(fz) > 0:
            return str(fz)
        bitmask = pos.get("zahn_bitmask") or 0
        if bitmask:
            from db import bitmask_to_fdi
            fdi_list = bitmask_to_fdi(int(bitmask))
            if fdi_list:
                return ",".join(str(z) for z in fdi_list[:6])
        return ""

    # ── Patientendaten ────────────────────────────────────────────────────────
    pat_name    = f"{patient_info.get('name','')}, {patient_info.get('vorname','')}"
    pat_kurz    = f"{patient_info.get('name','')} {patient_info.get('vorname','')}"
    kv_id       = kv_details.get("solid", "?")
    kv_datum    = kv_details.get("datum", "")
    kv_text     = kv_details.get("kurztext", "")
    today       = datetime.date.today().strftime("%d.%m.%Y")

    # Kosten: aus KI-Analyse (preferred) oder Charly-kv_details (fallback)
    if selected_positions and goz_prices:
        kva       = _build_kva(selected_positions, goz_prices)
        honorar   = kva["honorar"]
        material  = kva["material_est"]
        labor_total = kva["labor_brutto"]   # inkl. 7% MwSt
        gesamt    = kva["gesamt"]
        _use_kva  = True
    else:
        honorar     = float(kv_details.get("honorar") or 0)
        material    = float(kv_details.get("material") or 0)
        labor_total = float(kv_details.get("labor") or 0)
        gesamt      = honorar + material + labor_total
        kva         = None
        _use_kva    = False

    # ── SEITE 1: Anschreiben ─────────────────────────────────────────────────
    para(f"Herrn/Frau", size=10, space_after=0)
    para(pat_kurz, bold=True, size=11, space_after=0)
    para("", space_after=20)

    para(f"Pfungstadt, den {today}", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=20)

    # Betreff
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"Heil- und Kostenplan für die geplante Behandlung bei Ihnen selbst")
    r.bold = True; r.font.size = Pt(11)
    p2 = doc.add_paragraph(style="Normal")
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(f"Kostenvoranschlag: {kv_id}  |  {kv_text}")
    r2.font.size = Pt(10)

    para("Sehr geehrte Damen und Herren,", size=10, space_after=8)
    para(
        "anbei erhalten Sie unsere Kostenschätzung für die geplante zahnärztliche Behandlung. "
        "Die Abrechnung erfolgt nach der Gebührenordnung für Zahnärzte (GOZ 2012) sowie den "
        "gültigen Laborgebühren.",
        size=10, space_after=16
    )

    # Kostenzusammenfassung (Tabelle wie im PDF)
    para("Ihre voraussichtlichen Gesamtkosten:", bold=True, size=11, space_after=6)

    cost_tbl = doc.add_table(rows=5, cols=2)
    cost_tbl.style = 'Table Grid'
    rows_data = [
        ("Voraussichtliche Gesamtkosten", f"EUR {gesamt:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")),
        ("davon Honorar (GOZ)", f"EUR {honorar:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")),
        ("davon Materialkosten", f"EUR {material:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")),
        ("davon Laborkosten", f"EUR {labor_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")),
        ("Kostendatum", kv_datum),
    ]
    for i, (label, value) in enumerate(rows_data):
        r = cost_tbl.rows[i]
        r.cells[0].text = label
        r.cells[1].text = value
        r.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        r.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        if i == 0:
            r.cells[0].paragraphs[0].runs[0].bold = True
            r.cells[1].paragraphs[0].runs[0].bold = True
            set_cell_bg(r.cells[0], "D6E4F0")
            set_cell_bg(r.cells[1], "D6E4F0")

    para("", space_after=16)
    para(
        "Die Behandlung erfolgt in mehreren Sitzungen. Bitte wenden Sie sich bei Fragen jederzeit an uns.",
        size=10, space_after=20
    )
    para("Mit freundlichen Grüßen", size=10, space_after=30)
    para("Dr. med. dent. Christian Jung", bold=True, size=10)
    para("Dr. Jung Zahnklinik · City-Passage 1-6 · 64319 Pfungstadt", size=9)

    # ── ANLAGE 1: Honorar ────────────────────────────────────────────────────
    doc.add_page_break()

    # Anlage-Header
    hdr_tbl = doc.add_table(rows=1, cols=3)
    hdr_tbl.style = 'Table Grid'
    hdr_tbl.rows[0].cells[0].text = f"Name: {pat_kurz}"
    hdr_tbl.rows[0].cells[1].text = "Anlage 1"
    hdr_tbl.rows[0].cells[2].text = "Kostenvoranschlag Honorar"
    for cell in hdr_tbl.rows[0].cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(9)
        set_cell_bg(cell, "D6E4F0")
    doc.add_paragraph("").paragraph_format.space_after = Pt(6)

    # GOZ-Tabelle
    goz_tbl = doc.add_table(rows=1, cols=6)
    goz_tbl.style = 'Table Grid'
    for i, h in enumerate(["Zahn", "Anzahl", "Nr.", "Art der Leistung", "Satz", "EUR"]):
        c = goz_tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        set_cell_bg(c, "E8E8E8")

    col_widths = [Cm(1.5), Cm(1.2), Cm(1.5), Cm(8.5), Cm(1.5), Cm(2.0)]
    for i, w in enumerate(col_widths):
        for row in goz_tbl.rows:
            row.cells[i].width = w

    honorar_sum = 0.0

    if _use_kva and selected_positions:
        # KI-Analyse: zusammengefasst ODER zahnweise
        if gruppiert:
            rows_to_render = _group_kva_positions(kva["positionen"])
        else:
            # Zahnweise: Trennzeilen nach Zahn
            from collections import defaultdict as _dd
            by_zahn = _dd(list)
            for pos in kva["positionen"]:
                by_zahn[pos["zahn"]].append(pos)
            rows_to_render = None  # special handling below

        if gruppiert:
            for pos in rows_to_render:
                row = goz_tbl.add_row()
                hon = float(pos.get("honorar", 0))
                honorar_sum += hon
                p_info = goz_prices.get(pos["goz_nr"], {})
                leistung = p_info.get("goztext") or pos.get("text", "")
                vals = [
                    pos["zahn_str"],
                    str(pos.get("anzahl", 1)),
                    _goz_display(pos.get("goz_nr", "")),
                    leistung[:80],
                    f"{pos.get('faktor', 2.3):.1f}",
                    f"{hon:.2f}",
                ]
                for i, (cell, val) in enumerate(zip(row.cells, vals)):
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if i == 5:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            for zahn_nr in sorted(by_zahn.keys(), key=lambda x: int(x) if str(x).isdigit() else 0):
                zahn_sep = goz_tbl.add_row()
                cell_sep = zahn_sep.cells[0]
                for j in range(1, 6):
                    cell_sep = cell_sep.merge(zahn_sep.cells[j])
                cell_sep.text = f"Zahn {zahn_nr}"
                cell_sep.paragraphs[0].runs[0].bold = True
                cell_sep.paragraphs[0].runs[0].font.size = Pt(8)
                set_cell_bg(cell_sep, "F0F0F0")
                for pos in by_zahn[zahn_nr]:
                    row = goz_tbl.add_row()
                    hon = float(pos.get("honorar", 0))
                    honorar_sum += hon
                    p_info = goz_prices.get(pos["goz_nr"], {})
                    leistung = p_info.get("goztext") or pos.get("text", "")
                    vals = [
                        str(zahn_nr),
                        str(pos.get("anzahl", 1)),
                        _goz_display(pos.get("goz_nr", "")),
                        leistung[:80],
                        f"{pos.get('faktor', 2.3):.1f}",
                        f"{hon:.2f}",
                    ]
                    for i, (cell, val) in enumerate(zip(row.cells, vals)):
                        cell.text = val
                        cell.paragraphs[0].runs[0].font.size = Pt(8)
                        if i == 5:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        # Fallback: bestehende Charly-Positionen
        positionen = kv_details.get("positionen", []) if kv_details else []
        current_phase = None
        for pos in positionen:
            if not pos.get("goz_nr"):
                continue
            phase = pos.get("phase_bezeichnung", "")
            if phase != current_phase:
                current_phase = phase
                phase_row = goz_tbl.add_row()
                cell = phase_row.cells[0]
                for j in range(1, 6):
                    cell = cell.merge(phase_row.cells[j])
                cell.text = phase
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                set_cell_bg(cell, "F0F0F0")
            row = goz_tbl.add_row()
            betrag = float(pos.get("betrag") or 0)
            honorar_sum += betrag
            vals = [
                fdi_from_pos(pos),
                str(int(float(pos.get("anzahl") or 1))),
                _goz_display(pos.get("goz_nr", "")),
                (pos.get("goz_text") or "")[:80],
                str(pos.get("faktor", "")),
                f"{betrag:.2f}",
            ]
            for i, (cell, val) in enumerate(zip(row.cells, vals)):
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                if i == 5:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Summenzeile
    sum_row = goz_tbl.add_row()
    lbl = sum_row.cells[0].merge(sum_row.cells[4])
    lbl.text = "Voraussichtlicher Betrag Honorar EUR:"
    lbl.paragraphs[0].runs[0].bold = True
    lbl.paragraphs[0].runs[0].font.size = Pt(9)
    sum_row.cells[5].text = f"{honorar_sum:.2f}"
    sum_row.cells[5].paragraphs[0].runs[0].bold = True
    sum_row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)
    sum_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_bg(sum_row.cells[0], "D6E4F0")
    set_cell_bg(sum_row.cells[5], "D6E4F0")

    # ── ANLAGE 2: Material ───────────────────────────────────────────────────
    # KVA-Material: Schätzwerte aus GOZ-Katalog (schaetzbetrag), wenn KI-Analyse aktiv
    _kva_mat_rows = []
    if _use_kva and kva:
        for pos in kva["positionen"]:
            if float(pos.get("mat_est", 0)) > 0:
                _kva_mat_rows.append(pos)

    if _kva_mat_rows or kv_material:
        doc.add_page_break()
        hdr2 = doc.add_table(rows=1, cols=3)
        hdr2.style = 'Table Grid'
        hdr2.rows[0].cells[0].text = f"Name: {pat_kurz}"
        hdr2.rows[0].cells[1].text = "Anlage 2"
        hdr2.rows[0].cells[2].text = "Kostenvoranschlag Material"
        for cell in hdr2.rows[0].cells:
            run = cell.paragraphs[0].runs[0]
            run.bold = True; run.font.size = Pt(9)
            set_cell_bg(cell, "D6E4F0")
        doc.add_paragraph("").paragraph_format.space_after = Pt(6)

        mat_tbl = doc.add_table(rows=1, cols=4)
        mat_tbl.style = 'Table Grid'
        for i, h in enumerate(["Anzahl", "Nr.", "Material", "EUR"]):
            c = mat_tbl.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(8)
            set_cell_bg(c, "E8E8E8")

        mat_sum = 0.0
        # KI-Analyse: Schätzwerte aus GOZ-Katalog (ggf. zusammengefasst)
        items_to_show = []
        if _kva_mat_rows:
            rows_mat = _group_kva_positions(_kva_mat_rows) if gruppiert else _kva_mat_rows
            for pos in rows_mat:
                p_info = goz_prices.get(pos["goz_nr"], {})
                bezeichnung = p_info.get("goztext") or pos.get("text", "")
                zahn_s = pos.get("zahn_str") if gruppiert else f"Z{pos.get('zahn', '')}"
                items_to_show.append({
                    "anzahl": pos.get("anzahl", 1),
                    "nr": f"{zahn_s}/{pos['goz_nr']}",
                    "bez": f"Material: {bezeichnung[:50]}",
                    "betrag": float(pos.get("mat_est", 0)),
                })
        elif kv_material:
            for item in kv_material:
                items_to_show.append({
                    "anzahl": item.get("anzahl", ""),
                    "nr": str(item.get("mat_kuerzel") or item.get("mat_nr") or ""),
                    "bez": str(item.get("mat_bez") or "")[:70],
                    "betrag": float(item.get("betrag") or 0),
                })

        for item in items_to_show:
            betrag = float(item["betrag"])
            mat_sum += betrag
            row = mat_tbl.add_row()
            vals = [str(item["anzahl"]), item["nr"], item["bez"], f"{betrag:.2f}"]
            for i, (cell, val) in enumerate(zip(row.cells, vals)):
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                if i == 3:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        sum_row2 = mat_tbl.add_row()
        lbl2 = sum_row2.cells[0].merge(sum_row2.cells[2])
        lbl2.text = "Voraussichtlicher Betrag Material EUR:"
        lbl2.paragraphs[0].runs[0].bold = True
        lbl2.paragraphs[0].runs[0].font.size = Pt(9)
        sum_row2.cells[3].text = f"{mat_sum:.2f}"
        sum_row2.cells[3].paragraphs[0].runs[0].bold = True
        sum_row2.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        sum_row2.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_bg(sum_row2.cells[0], "D6E4F0")
        set_cell_bg(sum_row2.cells[3], "D6E4F0")

    # ── ANLAGE 3: Labor ──────────────────────────────────────────────────────
    lab_leistungen  = (kv_labor or {}).get("leistungen", [])
    lab_materialien = (kv_labor or {}).get("materialien", [])
    lab_summary     = (kv_labor or {}).get("summary", {})

    # KVA-Labor: Schätzwerte aus GOZ-Katalog (fremdschaetzbetr + fremdgoldbetr)
    _kva_lab_rows = []
    if _use_kva and kva:
        for pos in kva["positionen"]:
            if float(pos.get("lab_est", 0)) > 0:
                _kva_lab_rows.append(pos)

    if _kva_lab_rows or lab_leistungen or lab_materialien:
        doc.add_page_break()
        hdr3 = doc.add_table(rows=1, cols=3)
        hdr3.style = 'Table Grid'
        hdr3.rows[0].cells[0].text = f"Name: {pat_kurz}"
        hdr3.rows[0].cells[1].text = "Anlage 3"
        hdr3.rows[0].cells[2].text = "Kostenvoranschlag Labor"
        for cell in hdr3.rows[0].cells:
            run = cell.paragraphs[0].runs[0]
            run.bold = True; run.font.size = Pt(9)
            set_cell_bg(cell, "D6E4F0")
        para("Für die geplante Behandlung fallen folgende voraussichtliche Laborkosten an:", size=9, space_before=8, space_after=6)

        lab_tbl = doc.add_table(rows=1, cols=5)
        lab_tbl.style = 'Table Grid'
        for i, h in enumerate(["Anzahl", "Nr.", "Labor", "E/F", "EUR Netto"]):
            c = lab_tbl.rows[0].cells[i]
            c.text = h
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(8)
            set_cell_bg(c, "E8E8E8")

        lab_netto = 0.0
        if _kva_lab_rows:
            rows_lab = _group_kva_positions(_kva_lab_rows) if gruppiert else _kva_lab_rows
            for pos in rows_lab:
                betrag = float(pos.get("lab_est", 0))
                lab_netto += betrag
                p_info = goz_prices.get(pos["goz_nr"], {})
                bezeichnung = p_info.get("goztext") or pos.get("text", "")
                zahn_s = pos.get("zahn_str") if gruppiert else f"Z{pos.get('zahn', '')}"
                row = lab_tbl.add_row()
                vals = [
                    str(pos.get("anzahl", 1)),
                    f"{zahn_s}/{pos['goz_nr']}",
                    f"Fremdlabor: {bezeichnung[:52]}",
                    "F",
                    f"{betrag:.2f}",
                ]
                for i, (cell, val) in enumerate(zip(row.cells, vals)):
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if i == 4:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            for lp in lab_leistungen:
                betrag = float(lp.get("betrag") or 0)
                lab_netto += betrag
                row = lab_tbl.add_row()
                ef = "F" if lp.get("eigenfremd") == 1 else "E"
                vals = [
                    str(lp.get("anzahl", "")),
                    str(lp.get("nummer") or ""),
                    str(lp.get("bezeichnung") or "")[:65],
                    ef,
                    f"{betrag:.2f}",
                ]
                for i, (cell, val) in enumerate(zip(row.cells, vals)):
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if i == 4:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Labor-Material
        if lab_materialien:
            mat_sep = lab_tbl.add_row()
            lbl_mat = mat_sep.cells[0].merge(mat_sep.cells[4])
            lbl_mat.text = "Materialien:"
            lbl_mat.paragraphs[0].runs[0].bold = True
            lbl_mat.paragraphs[0].runs[0].font.size = Pt(8)
            set_cell_bg(lbl_mat, "F0F0F0")

            for lm in lab_materialien:
                betrag = float(lm.get("betrag") or 0)
                lab_netto += betrag
                row = lab_tbl.add_row()
                vals = [
                    str(lm.get("anzahl", "")),
                    str(lm.get("kuerzel") or lm.get("nummer") or ""),
                    str(lm.get("bezeichnung") or "")[:65],
                    "F",
                    f"{betrag:.2f}",
                ]
                for i, (cell, val) in enumerate(zip(row.cells, vals)):
                    cell.text = val
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                    if i == 4:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Labor-Summe mit MwSt (7% für Fremdlabor)
        fremd_lab  = float(lab_summary.get("fremd_labor") or 0) if lab_summary else lab_netto
        mwst_rate  = 0.07
        mwst_betrag = fremd_lab * mwst_rate
        lab_brutto  = fremd_lab + mwst_betrag

        sum_row3 = lab_tbl.add_row()
        lbl3 = sum_row3.cells[0].merge(sum_row3.cells[3])
        lbl3.text = f"Betrag Labor EUR {fremd_lab:.2f} + MwSt. EUR {mwst_betrag:.2f} ="
        lbl3.paragraphs[0].runs[0].bold = True
        lbl3.paragraphs[0].runs[0].font.size = Pt(9)
        sum_row3.cells[4].text = f"{lab_brutto:.2f}"
        sum_row3.cells[4].paragraphs[0].runs[0].bold = True
        sum_row3.cells[4].paragraphs[0].runs[0].font.size = Pt(9)
        sum_row3.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_bg(sum_row3.cells[0], "D6E4F0")
        set_cell_bg(sum_row3.cells[4], "D6E4F0")

    doc.save(str(output_path))


# ─────────────────────────────────────────────────────────────────────────────
# Projekte auflisten
# ─────────────────────────────────────────────────────────────────────────────

def list_hkp_projekte(limit: int = 20) -> list[dict]:
    """Listet gespeicherte HKP-Projekte auf (neueste zuerst)."""
    base = Path(PROJEKTE_PFAD) / "HKP"
    if not base.exists():
        return []
    projekte = []
    for meta_file in sorted(base.glob("*/projekt_meta.json"), reverse=True)[:limit]:
        try:
            with open(meta_file, encoding="utf-8") as f:
                meta = json.load(f)
                meta["ordner"] = str(meta_file.parent)
                projekte.append(meta)
        except Exception:
            pass
    return projekte
